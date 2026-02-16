# import os
# import json
# import pandas as pd
# import base64
# import xml.etree.ElementTree as ET
# from flask import Flask, request, jsonify, render_template_string
# from docx import Document
# from PyPDF2 import PdfReader

# app = Flask(__name__)

# stored_data = None


# # -----------------------------
# # Extract PDF
# # -----------------------------
# def extract_pdf_text(file):
#     reader = PdfReader(file)
#     text = ""
#     for page in reader.pages:
#         text += page.extract_text() or ""
#     return {"pdf_text": text}


# # -----------------------------
# # Extract DOCX
# # -----------------------------
# def extract_docx_text(file):
#     doc = Document(file)
#     text = "\n".join([p.text for p in doc.paragraphs])
#     return {"docx_text": text}


# # -----------------------------
# # XML to dict
# # -----------------------------
# def xml_to_dict(element):
#     return {
#         element.tag: {
#             "text": element.text,
#             "attributes": element.attrib,
#             "children": [xml_to_dict(child) for child in element]
#         }
#     }


# # -----------------------------
# # Professional Home Page
# # -----------------------------
# @app.route("/")
# def home():
#     return render_template_string("""
#     <!DOCTYPE html>
#     <html>
#     <head>
#         <title>Webhook Dashboard</title>
#         <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/css/bootstrap.min.css" rel="stylesheet">
#     </head>
#     <body class="bg-light">

#     <div class="container mt-5">
#         <div class="card shadow-lg p-4">
#             <h2 class="text-center mb-4">📂 Webhook File Dashboard</h2>

#             <div class="text-center">
#                 <a href="/upload" class="btn btn-primary btn-lg m-2">Upload File</a>
#                 <a href="/webhook" class="btn btn-success btn-lg m-2">View Stored JSON</a>
#             </div>
#         </div>
#     </div>

#     </body>
#     </html>
#     """)


# # -----------------------------
# # Upload Page
# # -----------------------------
# @app.route("/upload", methods=["GET", "POST"])
# def upload_file():
#     global stored_data

#     if request.method == "GET":
#         return render_template_string("""
#         <!DOCTYPE html>
#         <html>
#         <head>
#             <title>Upload File</title>
#             <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/css/bootstrap.min.css" rel="stylesheet">
#             <style>
#                 .drop-zone {
#                     border: 3px dashed #0d6efd;
#                     border-radius: 15px;
#                     padding: 50px;
#                     text-align: center;
#                     cursor: pointer;
#                     transition: 0.3s;
#                     background-color: #f8f9fa;
#                 }
#                 .drop-zone.dragover {
#                     background-color: #e7f1ff;
#                 }
#                 .file-name {
#                     margin-top: 15px;
#                     font-weight: bold;
#                     color: #198754;
#                 }
#             </style>
#         </head>
#         <body class="bg-light">

#         <div class="container mt-5">
#             <div class="card shadow-lg p-4">
#                 <h3 class="text-center mb-4">📂 Drag & Drop File Upload</h3>

#                 <form method="POST" enctype="multipart/form-data" id="uploadForm">
#                     <div class="drop-zone" id="dropZone">
#                         <p>Drag & Drop your file here</p>
#                         <p>or</p>
#                         <button type="button" class="btn btn-primary" onclick="document.getElementById('fileInput').click()">Browse File</button>
#                         <input type="file" name="file" id="fileInput" hidden required>
#                         <div class="file-name" id="fileName"></div>
#                     </div>

#                     <div class="text-center mt-4">
#                         <button type="submit" class="btn btn-success">Upload & Convert</button>
#                         <a href="/" class="btn btn-secondary">Back</a>
#                     </div>
#                 </form>
#             </div>
#         </div>

#         <script>
#             const dropZone = document.getElementById("dropZone");
#             const fileInput = document.getElementById("fileInput");
#             const fileName = document.getElementById("fileName");

#             dropZone.addEventListener("dragover", (e) => {
#                 e.preventDefault();
#                 dropZone.classList.add("dragover");
#             });

#             dropZone.addEventListener("dragleave", () => {
#                 dropZone.classList.remove("dragover");
#             });

#             dropZone.addEventListener("drop", (e) => {
#                 e.preventDefault();
#                 dropZone.classList.remove("dragover");

#                 const files = e.dataTransfer.files;
#                 fileInput.files = files;
#                 fileName.textContent = "Selected File: " + files[0].name;
#             });

#             fileInput.addEventListener("change", () => {
#                 fileName.textContent = "Selected File: " + fileInput.files[0].name;
#             });
#         </script>

#         </body>
#         </html>
#         """)

#     # -----------------------------
#     # POST Logic (Keep Same)
#     # -----------------------------
#     file = request.files["file"]
#     filename = file.filename.lower()

#     try:
#         if filename.endswith(".csv"):
#             df = pd.read_csv(file)
#             stored_data = df.to_dict(orient="records")

#         elif filename.endswith(".json"):
#             stored_data = json.load(file)

#         elif filename.endswith(".xlsx"):
#             df = pd.read_excel(file)
#             stored_data = df.to_dict(orient="records")

#         elif filename.endswith(".txt"):
#             stored_data = {"text_content": file.read().decode("utf-8", errors="ignore")}

#         elif filename.endswith(".xml"):
#             tree = ET.parse(file)
#             root = tree.getroot()
#             stored_data = xml_to_dict(root)

#         elif filename.endswith(".pdf"):
#             stored_data = extract_pdf_text(file)

#         elif filename.endswith(".docx"):
#             stored_data = extract_docx_text(file)

#         else:
#             stored_data = {
#                 "filename": filename,
#                 "base64_data": base64.b64encode(file.read()).decode("utf-8")
#             }

#         return render_template_string("""
#         <div style="text-align:center;margin-top:50px;">
#             <h3 style="color:green;">✅ File Converted Successfully</h3>
#             <a href="/webhook" style="font-size:18px;">View JSON</a><br><br>
#             <a href="/">Back to Dashboard</a>
#         </div>
#         """)

#     except Exception as e:
#         return f"Error: {str(e)}"

# # -----------------------------
# # Pretty JSON Viewer
# # -----------------------------
# @app.route("/webhook")
# def webhook():
#     global stored_data

#     if stored_data is None:
#         return jsonify({"message": "No data stored yet"})

#     formatted_json = json.dumps(stored_data, indent=4)

#     return render_template_string("""
#     <!DOCTYPE html>
#     <html>
#     <head>
#         <title>Stored JSON</title>
#         <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/css/bootstrap.min.css" rel="stylesheet">
#         <style>
#             pre {
#                 background-color: #1e1e1e;
#                 color: #00ff99;
#                 padding: 20px;
#                 border-radius: 8px;
#                 overflow-x: auto;
#             }
#         </style>
#     </head>
#     <body class="bg-light">

#     <div class="container mt-5">
#         <div class="card shadow-lg p-4">
#             <h3 class="mb-3">📄 Stored JSON Output</h3>
#             <pre>{{ data }}</pre>
#             <a href="/" class="btn btn-secondary mt-3">Back</a>
#         </div>
#     </div>

#     </body>
#     </html>
#     """, data=formatted_json)


# # -----------------------------
# # Run for Render
# # -----------------------------
# if __name__ == "__main__":
#     port = int(os.environ.get("PORT", 5000))
#     app.run(host="0.0.0.0", port=port)

import os
import json
import pandas as pd
import base64
import xml.etree.ElementTree as ET
from flask import Flask, request, jsonify, render_template_string
from docx import Document
from PyPDF2 import PdfReader

app = Flask(__name__)

stored_data = {}

# -----------------------------
# File Converters
# -----------------------------

def extract_pdf_text(file):
    reader = PdfReader(file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return {"pdf_text": text}

def extract_docx_text(file):
    doc = Document(file)
    return {"docx_text": "\n".join([p.text for p in doc.paragraphs])}

def xml_to_dict(element):
    return {
        element.tag: {
            "text": element.text,
            "attributes": element.attrib,
            "children": [xml_to_dict(child) for child in element]
        }
    }

# -----------------------------
# PROFESSIONAL DASHBOARD
# -----------------------------

@app.route("/")
def home():
    return render_template_string("""
<!DOCTYPE html>
<html>
<head>
<title>Professional Webhook Dashboard</title>
<meta name="viewport" content="width=device-width, initial-scale=1">
<link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/css/bootstrap.min.css" rel="stylesheet">
<style>
body { background: linear-gradient(135deg, #1f4037, #99f2c8); min-height:100vh; }
.card { border-radius: 15px; }
.drop-zone {
    border: 3px dashed #0d6efd;
    border-radius: 15px;
    padding: 50px;
    text-align: center;
    background: #f8f9fa;
    transition: 0.3s;
}
.drop-zone.dragover {
    background: #e7f1ff;
}
.file-list {
    margin-top: 15px;
}
pre {
    background: #111;
    color: #00ff9d;
    padding: 20px;
    border-radius: 10px;
    max-height: 400px;
    overflow:auto;
}
</style>
</head>
<body>

<div class="container py-5">
<div class="card shadow-lg p-4">

<h2 class="text-center mb-4">📂 Multi File Webhook Dashboard</h2>

<form method="POST" action="/upload" enctype="multipart/form-data" id="uploadForm">

<div class="drop-zone" id="dropZone">
    <h5>Drag & Drop Files Here</h5>
    <p>or</p>
    <button type="button" class="btn btn-primary" onclick="fileInput.click()">Browse Files</button>
    <input type="file" id="fileInput" name="file" multiple hidden required>
    <div class="file-list" id="fileList"></div>
</div>

<div class="text-center mt-4">
    <button type="submit" class="btn btn-success">Upload & Convert</button>
    <a href="/webhook" class="btn btn-dark">View JSON</a>
</div>

</form>

</div>
</div>

<script>
const dropZone = document.getElementById("dropZone");
const fileInput = document.getElementById("fileInput");
const fileList = document.getElementById("fileList");

dropZone.addEventListener("dragover", e => {
    e.preventDefault();
    dropZone.classList.add("dragover");
});

dropZone.addEventListener("dragleave", () => {
    dropZone.classList.remove("dragover");
});

dropZone.addEventListener("drop", e => {
    e.preventDefault();
    dropZone.classList.remove("dragover");
    fileInput.files = e.dataTransfer.files;
    updateFileList();
});

fileInput.addEventListener("change", updateFileList);

function updateFileList() {
    fileList.innerHTML = "";
    Array.from(fileInput.files).forEach(file => {
        fileList.innerHTML += "<div>📄 " + file.name + "</div>";
    });
}
</script>

</body>
</html>
""")

# -----------------------------
# Upload Logic
# -----------------------------

@app.route("/upload", methods=["POST"])
def upload():
    global stored_data

    files = request.files.getlist("file")

    for file in files:
        filename = file.filename.lower()

        try:
            if filename.endswith(".csv"):
                df = pd.read_csv(file)
                stored_data[filename] = df.to_dict(orient="records")

            elif filename.endswith(".json"):
                stored_data[filename] = json.load(file)

            elif filename.endswith(".xlsx"):
                df = pd.read_excel(file)
                stored_data[filename] = df.to_dict(orient="records")

            elif filename.endswith(".txt"):
                stored_data[filename] = {"text": file.read().decode("utf-8", errors="ignore")}

            elif filename.endswith(".xml"):
                tree = ET.parse(file)
                stored_data[filename] = xml_to_dict(tree.getroot())

            elif filename.endswith(".pdf"):
                stored_data[filename] = extract_pdf_text(file)

            elif filename.endswith(".docx"):
                stored_data[filename] = extract_docx_text(file)

            else:
                stored_data[filename] = {
                    "base64_data": base64.b64encode(file.read()).decode("utf-8")
                }

        except Exception as e:
            stored_data[filename] = {"error": str(e)}

    return """
    <div style='text-align:center;margin-top:50px;'>
    <h3 style='color:green;'>✅ Files Converted Successfully</h3>
    <a href='/'>Back</a> | <a href='/webhook'>View JSON</a>
    </div>
    """

# -----------------------------
# JSON Viewer
# -----------------------------

@app.route("/webhook")
def webhook():
    if not stored_data:
        return "<h3 style='text-align:center;margin-top:50px;'>No files uploaded yet.</h3>"

    formatted = json.dumps(stored_data, indent=4)

    return f"""
    <div style='padding:40px;background:#111;color:#00ff9d;font-family:monospace;'>
    <h2>Stored JSON Output</h2>
    <pre>{formatted}</pre>
    <a href='/' style='color:white;'>Back</a>
    </div>
    """

# -----------------------------
# Run
# -----------------------------

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)

